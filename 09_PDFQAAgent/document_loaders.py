from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple
import io

import requests
from pypdf import PdfReader
from docx import Document as DocxDocument
import trafilatura
from bs4 import BeautifulSoup


@dataclass
class LoadedDocument:
    title: str
    source: str
    pages: List[str]
    media_type: str  # pdf|docx|txt|html


def _clean_text(text: str) -> str:
    # Normalize whitespace and remove very short noisy lines
    lines = [l.strip() for l in text.splitlines()]
    lines = [l for l in lines if l and len(l) > 2]
    return "\n".join(lines)


def load_pdf_from_path(path: Path) -> LoadedDocument:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(_clean_text(page.extract_text() or ''))
    title = Path(path).stem
    return LoadedDocument(title=title, source=str(path), pages=pages, media_type='pdf')


def load_docx_from_path(path: Path) -> LoadedDocument:
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    text = _clean_text('\n'.join(paragraphs))
    return LoadedDocument(title=Path(path).stem, source=str(path), pages=[text], media_type='docx')


def load_txt_from_path(path: Path) -> LoadedDocument:
    text = Path(path).read_text(encoding='utf-8', errors='ignore')
    return LoadedDocument(title=Path(path).stem, source=str(path), pages=[_clean_text(text)], media_type='txt')


def fetch_url(url: str) -> Tuple[str, bytes, Optional[str]]:
    resp = requests.get(url, timeout=30)
    resp.raise_for_status()
    content_type = resp.headers.get('Content-Type', '').lower()
    filename = url.split('/')[-1] or 'downloaded'
    return filename, resp.content, content_type


def _extract_html_bs4(content: bytes) -> str:
    soup = BeautifulSoup(content, 'lxml')
    # Remove script/style/nav/footer
    for tag in soup(['script', 'style', 'noscript']):
        tag.decompose()
    text = soup.get_text(separator='\n')
    return _clean_text(text)


def load_from_url(url: str) -> LoadedDocument:
    filename, content, content_type = fetch_url(url)
    if (content_type and 'application/pdf' in content_type) or filename.lower().endswith('.pdf'):
        reader = PdfReader(io.BytesIO(content))
        pages = [ _clean_text(page.extract_text() or '') for page in reader.pages ]
        return LoadedDocument(title=filename.rsplit('.', 1)[0], source=url, pages=pages, media_type='pdf')
    if filename.lower().endswith('.docx'):
        fp = io.BytesIO(content)
        doc = DocxDocument(fp)
        paragraphs = [p.text for p in doc.paragraphs]
        return LoadedDocument(title=filename.rsplit('.', 1)[0], source=url, pages=['\n'.join(paragraphs)], media_type='docx')
    # HTML extraction
    downloaded = trafilatura.extract(content, include_comments=False, include_tables=False)
    text = downloaded or ''
    if not text or len(text.strip()) < 40:
        text = _extract_html_bs4(content)
    text = _clean_text(text)
    return LoadedDocument(title=filename, source=url, pages=[text], media_type='html')