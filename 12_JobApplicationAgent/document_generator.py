"""
Document generation utilities for creating PDF and DOCX files
"""

import os
import io
import logging
from typing import Dict, List, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
except ImportError:
    SimpleDocTemplate = None

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Handles document generation for resumes and cover letters"""
    
    def __init__(self):
        self.supported_formats = ["pdf", "docx"]
    
    def generate_resume_document(
        self,
        content: str,
        format_type: str = "pdf",
        filename: Optional[str] = None
    ) -> bytes:
        """Generate resume document in specified format"""
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}")
        
        if format_type == "pdf":
            return self._generate_pdf_resume(content, filename)
        elif format_type == "docx":
            return self._generate_docx_resume(content, filename)
    
    def generate_cover_letter_document(
        self,
        content: str,
        format_type: str = "pdf",
        filename: Optional[str] = None
    ) -> bytes:
        """Generate cover letter document in specified format"""
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}")
        
        if format_type == "pdf":
            return self._generate_pdf_cover_letter(content, filename)
        elif format_type == "docx":
            return self._generate_docx_cover_letter(content, filename)
    
    def generate_additional_document(
        self,
        content: str,
        doc_type: str,
        format_type: str = "pdf",
        filename: Optional[str] = None
    ) -> bytes:
        """Generate additional document in specified format"""
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}")
        
        if format_type == "pdf":
            return self._generate_pdf_additional_document(content, doc_type, filename)
        elif format_type == "docx":
            return self._generate_docx_additional_document(content, doc_type, filename)
    
    def _generate_pdf_resume(self, content: str, filename: Optional[str] = None) -> bytes:
        """Generate PDF resume with improved formatting"""
        if not SimpleDocTemplate:
            raise ImportError("ReportLab is required for PDF generation")
        
        try:
            # Create buffer for PDF
            buffer = io.BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles with improved formatting
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12,
                alignment=TA_CENTER,
                textColor=colors.darkblue,
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=8,
                spaceBefore=16,
                textColor=colors.darkblue,
                fontName='Helvetica-Bold',
                leftIndent=0
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                leading=14,
                fontName='Helvetica',
                leftIndent=0
            )
            
            bullet_style = ParagraphStyle(
                'CustomBullet',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=4,
                leading=14,
                fontName='Helvetica',
                leftIndent=20,
                bulletIndent=10
            )
            
            # Build content
            story = []
            
            # Parse content and create PDF elements with better structure
            lines = content.split('\n')
            current_section = ""
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if this is a section header
                if self._is_section_header(line):
                    current_section = line
                    story.append(Paragraph(line.upper(), heading_style))
                    story.append(Spacer(1, 8))
                elif line.startswith(('•', '-', '*', '○', '▪', '▫')):
                    # Bullet points
                    clean_line = line.lstrip('•-*○▪▫ ').strip()
                    if clean_line:
                        story.append(Paragraph(f"• {clean_line}", bullet_style))
                elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                    # Numbered lists
                    story.append(Paragraph(line, normal_style))
                else:
                    # Regular content
                    story.append(Paragraph(line, normal_style))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"PDF resume generation failed: {e}")
            raise ValueError("Failed to generate PDF resume")
    
    def _generate_docx_resume(self, content: str, filename: Optional[str] = None) -> bytes:
        """Generate DOCX resume with improved formatting"""
        if not Document:
            raise ImportError("python-docx is required for DOCX generation")
        
        try:
            # Create document
            doc = Document()
            
            # Set up styles
            self._setup_docx_styles(doc)
            
            # Parse content and add to document with better structure
            lines = content.split('\n')
            current_section = ""
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if this is a section header
                if self._is_section_header(line):
                    current_section = line
                    heading = doc.add_heading(line.upper(), level=2)
                    heading.style = doc.styles['Heading 2']
                elif line.startswith(('•', '-', '*', '○', '▪', '▫')):
                    # Bullet points
                    clean_line = line.lstrip('•-*○▪▫ ').strip()
                    if clean_line:
                        paragraph = doc.add_paragraph()
                        paragraph.style = doc.styles['List Bullet']
                        paragraph.add_run(clean_line)
                elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                    # Numbered lists
                    paragraph = doc.add_paragraph()
                    paragraph.style = doc.styles['List Number']
                    paragraph.add_run(line)
                else:
                    # Regular content
                    paragraph = doc.add_paragraph(line)
                    paragraph.style = doc.styles['Normal']
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            
            # Get DOCX bytes
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            return docx_bytes
            
        except Exception as e:
            logger.error(f"DOCX resume generation failed: {e}")
            raise ValueError("Failed to generate DOCX resume")
    
    def _generate_pdf_cover_letter(self, content: str, filename: Optional[str] = None) -> bytes:
        """Generate PDF cover letter with improved formatting"""
        if not SimpleDocTemplate:
            raise ImportError("ReportLab is required for PDF generation")
        
        try:
            # Create buffer for PDF
            buffer = io.BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=1*inch,
                leftMargin=1*inch,
                topMargin=1*inch,
                bottomMargin=1*inch
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles for cover letter
            title_style = ParagraphStyle(
                'CoverLetterTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=20,
                alignment=TA_CENTER,
                textColor=colors.darkblue,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CoverLetterNormal',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leading=16,
                alignment=TA_LEFT,
                fontName='Helvetica',
                firstLineIndent=0
            )
            
            # Build content
            story = []
            
            # Add title
            story.append(Paragraph("Cover Letter", title_style))
            story.append(Spacer(1, 20))
            
            # Add date
            current_date = datetime.now().strftime("%B %d, %Y")
            story.append(Paragraph(current_date, normal_style))
            story.append(Spacer(1, 20))
            
            # Add content with proper paragraph formatting
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Clean up the paragraph
                    clean_para = paragraph.strip().replace('\n', ' ')
                    story.append(Paragraph(clean_para, normal_style))
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"PDF cover letter generation failed: {e}")
            raise ValueError("Failed to generate PDF cover letter")
    
    def _generate_docx_cover_letter(self, content: str, filename: Optional[str] = None) -> bytes:
        """Generate DOCX cover letter with improved formatting"""
        if not Document:
            raise ImportError("python-docx is required for DOCX generation")
        
        try:
            # Create document
            doc = Document()
            
            # Set up styles
            self._setup_docx_styles(doc)
            
            # Add title
            title = doc.add_heading("Cover Letter", 0)
            title.style = doc.styles['Title']
            
            # Add date
            current_date = datetime.now().strftime("%B %d, %Y")
            date_paragraph = doc.add_paragraph(current_date)
            date_paragraph.style = doc.styles['Normal']
            
            # Add spacing
            doc.add_paragraph()
            
            # Add content with proper paragraph formatting
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Clean up the paragraph
                    clean_para = paragraph.strip().replace('\n', ' ')
                    p = doc.add_paragraph(clean_para)
                    p.style = doc.styles['Normal']
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            
            # Get DOCX bytes
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            return docx_bytes
            
        except Exception as e:
            logger.error(f"DOCX cover letter generation failed: {e}")
            raise ValueError("Failed to generate DOCX cover letter")
    
    def _generate_pdf_additional_document(self, content: str, doc_type: str, filename: Optional[str] = None) -> bytes:
        """Generate PDF additional document with improved formatting"""
        if not SimpleDocTemplate:
            raise ImportError("ReportLab is required for PDF generation")
        
        try:
            # Create buffer for PDF
            buffer = io.BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=1*inch,
                leftMargin=1*inch,
                topMargin=1*inch,
                bottomMargin=1*inch
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'DocumentTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=20,
                alignment=TA_CENTER,
                textColor=colors.darkblue,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'DocumentNormal',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leading=16,
                alignment=TA_LEFT,
                fontName='Helvetica',
                firstLineIndent=0
            )
            
            # Build content
            story = []
            
            # Add title based on document type
            doc_titles = {
                'personal_statement': 'Personal Statement',
                'reference_page': 'Professional References',
                'thank_you_note': 'Thank You Note',
                'motivation_letter': 'Motivation Letter',
                'linkedin_bio': 'LinkedIn Bio'
            }
            title = doc_titles.get(doc_type, 'Document')
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            # Add content with proper formatting
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    clean_para = paragraph.strip().replace('\n', ' ')
                    story.append(Paragraph(clean_para, normal_style))
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"PDF {doc_type} generation failed: {e}")
            raise ValueError(f"Failed to generate PDF {doc_type}")
    
    def _generate_docx_additional_document(self, content: str, doc_type: str, filename: Optional[str] = None) -> bytes:
        """Generate DOCX additional document with improved formatting"""
        if not Document:
            raise ImportError("python-docx is required for DOCX generation")
        
        try:
            # Create document
            doc = Document()
            
            # Set up styles
            self._setup_docx_styles(doc)
            
            # Add title based on document type
            doc_titles = {
                'personal_statement': 'Personal Statement',
                'reference_page': 'Professional References',
                'thank_you_note': 'Thank You Note',
                'motivation_letter': 'Motivation Letter',
                'linkedin_bio': 'LinkedIn Bio'
            }
            title = doc_titles.get(doc_type, 'Document')
            title_heading = doc.add_heading(title, 0)
            title_heading.style = doc.styles['Title']
            
            # Add spacing
            doc.add_paragraph()
            
            # Add content with proper formatting
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    clean_para = paragraph.strip().replace('\n', ' ')
                    p = doc.add_paragraph(clean_para)
                    p.style = doc.styles['Normal']
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            
            # Get DOCX bytes
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            return docx_bytes
            
        except Exception as e:
            logger.error(f"DOCX {doc_type} generation failed: {e}")
            raise ValueError(f"Failed to generate DOCX {doc_type}")
    
    def _setup_docx_styles(self, doc: Document):
        """Set up custom styles for DOCX document with improved formatting"""
        try:
            # Title style
            title_style = doc.styles['Title']
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(18)
            title_style.font.bold = True
            title_style.font.color.rgb = None  # Default color
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            
            # Heading 1 style
            heading1_style = doc.styles['Heading 1']
            heading1_style.font.name = 'Calibri'
            heading1_style.font.size = Pt(16)
            heading1_style.font.bold = True
            heading1_style.font.color.rgb = None
            heading1_style.paragraph_format.space_after = Pt(8)
            
            # Heading 2 style
            heading2_style = doc.styles['Heading 2']
            heading2_style.font.name = 'Calibri'
            heading2_style.font.size = Pt(14)
            heading2_style.font.bold = True
            heading2_style.font.color.rgb = None
            heading2_style.paragraph_format.space_after = Pt(6)
            heading2_style.paragraph_format.space_before = Pt(12)
            
            # Normal style
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.15
            
            # List Bullet style
            if 'List Bullet' in doc.styles:
                bullet_style = doc.styles['List Bullet']
                bullet_style.font.name = 'Calibri'
                bullet_style.font.size = Pt(11)
                bullet_style.paragraph_format.space_after = Pt(4)
                bullet_style.paragraph_format.left_indent = Pt(20)
            
            # List Number style
            if 'List Number' in doc.styles:
                number_style = doc.styles['List Number']
                number_style.font.name = 'Calibri'
                number_style.font.size = Pt(11)
                number_style.paragraph_format.space_after = Pt(4)
                number_style.paragraph_format.left_indent = Pt(20)
            
        except Exception as e:
            logger.warning(f"Could not set up custom styles: {e}")
    
    def _is_section_header(self, line: str) -> bool:
        """Check if a line is a section header with improved detection"""
        # Common section headers
        section_headers = [
            "experience", "work experience", "employment history", "professional experience",
            "education", "academic background", "qualifications", "academic experience",
            "skills", "technical skills", "competencies", "key skills", "core competencies",
            "summary", "objective", "profile", "about", "professional summary",
            "contact", "personal information", "contact information",
            "projects", "project experience", "portfolio",
            "certifications", "certificates", "licenses",
            "awards", "achievements", "honors",
            "publications", "research", "papers",
            "volunteer", "volunteer experience", "community service",
            "languages", "language skills", "foreign languages"
        ]
        
        line_lower = line.lower().strip()
        
        # Check if line matches any section header
        for header in section_headers:
            if header in line_lower and len(line) < 50:  # Section headers are usually short
                return True
        
        # Check for all caps (common for section headers)
        if line.isupper() and len(line) < 30:
            return True
        
        # Check for patterns like "SECTION NAME" or "Section Name"
        if line.count(' ') <= 3 and any(word.isupper() for word in line.split()):
            return True
        
        # Check for common formatting patterns
        if line.endswith(':') and len(line) < 30:
            return True
        
        return False
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats"""
        return self.supported_formats.copy()
    
    def validate_format(self, format_type: str) -> bool:
        """Validate if format is supported"""
        return format_type in self.supported_formats
